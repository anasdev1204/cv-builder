from pathlib import Path

import pytest
from docx import Document

from models.cv import (
    CV,
    SectionMeta,
    UserData,
    Address,
    SectionEntry,
)

from services.cv_compiler import CVCompiler


@pytest.fixture
def sample_cv():
    return CV(
        user_data=UserData(
            name="John Doe",
            email="john@example.com",
            picture="./test_output/cv_pic.png",
            phone_number="+44 123456789",
            linkedin="linkedin.com/in/johndoe",
            portfolio="johndoe.dev",
            address=Address(
                city="London",
                country="UK",
            ),
        ),
        sections={
            "en": {
                "summary": {
                    "title": "Professional Summary",
                    "content": (
                        "Software engineer with experience building"
                        "scalable web applications and APIs."
                        "Looking for opportunities to contribute to innovative projects."
                    )
                } ,
                "experience": {
                    "title": "Professional Experience",
                    "content": [
                        SectionEntry(
                            title="Software Engineer",
                            subtitle="Company A",
                            start_date="2024",
                            end_date="2026",
                            bullet_points=[
                            "Developed Python services.",
                            "Built REST APIs using FastAPI.",
                            "Improved system performance by 35%.",
                            ],
                        ),
                        SectionEntry(
                            title="Software Developer Intern",
                            subtitle="Company B",
                            start_date="2023",
                            end_date="2024",
                            bullet_points=[
                                "Implemented frontend features.",
                                "Worked with React and TypeScript.",
                            ],
                        ),
                    ]
                },
                "education": {
                    "title": "Education",
                    "content": [
                        SectionEntry(
                            title="BSc Computer Science",
                            subtitle="University of Example",
                            start_date="2021",
                            end_date="2025",
                            bullet_points=[
                                "Graduated with First Class Honours."
                            ],
                        )
                    ]
                },
                "languages": {
                    "title": "Languages",
                    "content": ["English", "French"]
                },
                "skills": {
                    "title": "Skills",
                    "content": ["Python", "FastAPI", "React", "TypeScript"]
                },
                "other_sections": {
                    "Certifications": {
                        "title": "Certifications",
                        "content": [
                            SectionEntry(
                                title="AWS Certified Developer",
                                subtitle="Amazon Web Services",
                                start_date="2025",
                                bullet_points=[],
                            )
                        ]
                    }
                }
            }
        },
    )

@pytest.fixture
def compiler():
    return CVCompiler(
        output_directory="./test_output"
    )

def get_all_paragraph_text(document):
    return "\n".join(
        paragraph.text
        for paragraph in document.paragraphs
    )

@pytest.mark.asyncio
async def test_compile_docx_creates_file(
    compiler,
    sample_cv,
):
    output_path = await compiler.compile(
        cv=sample_cv,
        output_format="docx",
        selected_version="en"
    )

    path = Path(output_path)

    assert path.exists()
    assert path.is_file()
    assert path.suffix == ".docx"

@pytest.mark.asyncio
async def test_compile_docx_contains_personal_information(
    compiler,
    sample_cv,
):
    output_path = await compiler.compile(
        cv=sample_cv,
        output_format="docx",
    )

    document = Document(output_path)
    text = get_all_paragraph_text(document)

    assert "John Doe" in text
    assert "john@example.com" in text
    assert "+44 123456789" in text
    assert "London, UK" in text
    assert "linkedin.com/in/johndoe" in text
    assert "johndoe.dev" in text


@pytest.mark.asyncio
async def test_compile_docx_contains_summary(
    compiler,
    sample_cv: CV,
):
    output_path = await compiler.compile(
        cv=sample_cv,
        output_format="docx",
    )

    document = Document(output_path)
    text = get_all_paragraph_text(document)

    data: SectionMeta = sample_cv.sections["en"].summary

    assert data.title in text
    assert data.content in text


@pytest.mark.asyncio
async def test_compile_docx_contains_experience(
    compiler,
    sample_cv,
):
    output_path = await compiler.compile(
        cv=sample_cv,
        output_format="docx",
    )

    document = Document(output_path)
    text = get_all_paragraph_text(document)

    data: SectionMeta = sample_cv.sections["en"].experience

    assert data.title in text

    experiences = data.content
    for e in experiences:
        assert e.title in text
        assert e.subtitle in text
        assert e.start_date in text
        assert e.end_date in text
        assert all(bp in text for bp in e.bullet_points)


@pytest.mark.asyncio
async def test_compile_docx_contains_education(
    compiler,
    sample_cv,
):
    output_path = await compiler.compile(
        cv=sample_cv,
        output_format="docx",
    )

    document = Document(output_path)
    text = get_all_paragraph_text(document)

    data: SectionMeta = sample_cv.sections["en"].education

    assert data.title in text

    educations = data.content
    for e in educations:
        assert e.title in text
        assert e.subtitle in text
        assert e.start_date in text
        assert e.end_date in text
        assert all(bp in text for bp in e.bullet_points)


@pytest.mark.asyncio
async def test_compile_docx_contains_skills(
    compiler,
    sample_cv,
):
    output_path = await compiler.compile(
        cv=sample_cv,
        output_format="docx",
    )

    document = Document(output_path)
    text = get_all_paragraph_text(document)

    data: SectionMeta = sample_cv.sections["en"].skills

    assert data.title in text

    for skill in data.content:
        assert skill in text


@pytest.mark.asyncio
async def test_compile_docx_contains_languages(
    compiler,
    sample_cv,
):
    output_path = await compiler.compile(
        cv=sample_cv,
        output_format="docx",
    )

    document = Document(output_path)
    text = get_all_paragraph_text(document)

    data: SectionMeta = sample_cv.sections["en"].languages

    assert data.title in text
    assert all(lang in text for lang in data.content)


@pytest.mark.asyncio
async def test_compile_docx_contains_other_sections(
    compiler,
    sample_cv,
):
    output_path = await compiler.compile(
        cv=sample_cv,
        output_format="docx",
    )

    document = Document(output_path)
    text = get_all_paragraph_text(document)

    data_dict: dict = sample_cv.sections["en"].other_sections

    for _, data in data_dict.items():
        data: SectionMeta = data
        assert data.title in text

        if isinstance(data.content, str):
            assert data.content in text
        elif isinstance(data.content, list):
            for entry in data.content:

                if isinstance(entry, str):
                    assert entry in text
                elif isinstance(entry, SectionEntry):
                    assert entry.title in text
                    assert entry.subtitle in text
                    assert entry.start_date in text
                    assert all(bp in text for bp in entry.bullet_points)

@pytest.mark.asyncio
async def test_filename_is_based_on_user_name(
    compiler,
    sample_cv,
):
    output_path = await compiler.compile(
        cv=sample_cv,
        output_format="docx",
    )

    assert Path(output_path).name == "John Doe.docx"


def test_safe_filename_removes_invalid_characters():
    result = CVCompiler._safe_filename(
        "John / Doe: Test?"
    )

    assert result == "John  Doe Test"


@pytest.mark.asyncio
async def test_unsupported_format_raises_error(
    compiler,
    sample_cv,
):
    with pytest.raises(
        ValueError,
        match="Unsupported format"
    ):
        await compiler.compile(
            cv=sample_cv,
            output_format="txt",
        )


@pytest.mark.asyncio
async def test_pdf_conversion_is_called(
    compiler,
    sample_cv,
    monkeypatch,
    tmp_path,
):
    docx_path = tmp_path / "John Doe.docx"
    docx_path.touch()

    called = {}

    def fake_convert(path):
        called["path"] = path
        return str(tmp_path / "John Doe.pdf")

    monkeypatch.setattr(
        compiler,
        "_convert_to_pdf",
        fake_convert,
    )

    result = await compiler.compile(
        cv=sample_cv,
        output_format="pdf",
    )

    assert called["path"].suffix == ".docx"
    assert result.endswith(".pdf")
