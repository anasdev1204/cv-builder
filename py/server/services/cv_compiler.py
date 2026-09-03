from pathlib import Path
import subprocess
from turtle import st
from docx.oxml import OxmlElement
from fastapi.responses import FileResponse

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn

from models.cv import CV, CvSections, SectionEntry, SectionMeta, UserData

from pathlib import Path

import yaml

from models.template import (
    BulletConfig,
    DateConfig,
    FontConfig,
    TemplateConfig,
    SectionRendererConfig,
    EntryConfig
)


class TemplateLoader:

    def __init__(self, config_directory: str = "configs"):
        self.config_directory = Path(config_directory)

    def load(self, template: str) -> TemplateConfig:

        path = self.config_directory / f"{template}.yaml"

        if not path.exists():
            raise ValueError(
                f"Template '{template}' not found: {path}"
            )

        with path.open("r", encoding="utf-8") as file:
            data = yaml.safe_load(file) or {}

        try:
            return TemplateConfig.model_validate(data)
        except Exception as e:
            raise ValueError(
                f"Invalid configuration for template '{template}': {e}"
            ) from e


class CVCompiler:
    def __init__(
        self,
        output_directory: str = "test_output",
        config_directory: str = "configs",
        template_config: TemplateConfig | None = None,
    ):
        self.output_directory = Path(output_directory)
        self.output_directory.mkdir(
            parents=True,
            exist_ok=True
        )

        self.template_loader = TemplateLoader(
            config_directory
        )

        self.template_config = template_config

    async def compile(
        self,
        cv: CV,
        template: str = "professional",
        output_format: str = "pdf",
        selected_version: str = "en",
        job_title: str = "Unknown",
        local: bool = True,
    ):
        docx_path = self._compile_docx(
            cv=cv,
            job_title=job_title,
            template=template,
            selected_version=selected_version,
        )

        if output_format == "docx":
            output_path = docx_path

        elif output_format == "pdf":
            output_path = self._convert_to_pdf(docx_path)

        else:
            raise ValueError(
                f"Unsupported format: {output_format}"
            )

        if local:
            return str(output_path)

        return self._compile_doc_url(output_path)

    def _compile_docx(
        self,
        cv: CV,
        job_title: str,
        template: str,
        selected_version: str,
    ) -> Path:
        
        if self.template_config is not None:
            config = self.template_config
        else:
            config = self.template_loader.load(template)

        document = Document()

        user_data = cv.user_data

        if selected_version not in cv.sections:
            raise ValueError(
                f"Selected version '{selected_version}' not found."
            )

        cv_version = cv.sections[selected_version]

        self._configure_page(
            document,
            config,
        )

        self._configure_styles(
            document,
            config,
        )

        self._add_header(
            document,
            user_data,
            job_title,
            config,
        )

        self._render_all_sections(
            document,
            cv_version,
            config,
        )

        output_path = (
            self.output_directory
            / f"{self._safe_filename(user_data.name)}.docx"
        )

        document.save(output_path)

        return output_path

    def _configure_page(
        self,
        document: Document,
        config: TemplateConfig,
    ):

        margins = config.page.margins

        section = document.sections[0]

        section.top_margin = Inches(margins.top)
        section.bottom_margin = Inches(margins.bottom)
        section.left_margin = Inches(margins.left)
        section.right_margin = Inches(margins.right)

    def _configure_styles(
        self,
        document: Document,
        config: TemplateConfig,
    ):

        normal = document.styles["Normal"]

        normal.font.name = config.font.family
        normal.font.size = Pt(config.font.size)

        normal.font.bold = config.font.bold
        normal.font.italic = config.font.italic
        normal.font.underline = config.font.underline
        normal.font.color.rgb = RGBColor.from_string(
            config.font.color.lstrip("#")
        )
        normal.font.character_spacing = Pt(config.font.character_spacing
        )

    def _add_header(
        self,
        document: Document,
        user_data: UserData,
        job_title: str,
        config: TemplateConfig,
    ):
        alignment_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }

        header_config = config.header
        alignment = alignment_map[header_config.alignment]

        if user_data.picture and header_config.show_picture:
            try:
                document.add_picture(
                    user_data.picture,
                    width=Inches(header_config.picture_size),
                )
                document.paragraphs[-1].alignment = alignment
            except Exception as e:
                print(f"Error adding picture: {e}")

        paragraph = document.add_paragraph()
        paragraph.alignment = alignment
        paragraph.paragraph_format.space_before = Pt(
            header_config.space_before
        )
        paragraph.paragraph_format.space_after = Pt(0)

        self._apply_font(
            paragraph.add_run(user_data.name),
            header_config.name,
        )

        if job_title:
            print(f"Adding job title: {job_title}")
            paragraph = document.add_paragraph()
            paragraph.alignment = alignment
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)

            self._apply_font(
                paragraph.add_run(job_title),
                header_config.job_title,
            )

        contact = []

        if user_data.email:
            contact.append(user_data.email)

        if user_data.phone_number:
            contact.append(user_data.phone_number)

        if user_data.address:
            contact.append(
                f"{user_data.address.city}, {user_data.address.country}"
            )

        if user_data.linkedin:
            contact.append(user_data.linkedin)

        if user_data.portfolio:
            contact.append(user_data.portfolio)

        if contact:
            paragraph = document.add_paragraph()
            paragraph.alignment = alignment
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(
                header_config.space_after
            )

            self._apply_font(
                paragraph.add_run(
                    header_config.contact.separator.join(contact)
                ),
                header_config.contact,
            )

        if header_config.show_divider:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)

            p = paragraph._p
            p_pr = p.get_or_add_pPr()

            p_bdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")

            bottom.set(qn("w:val"), "single")
            bottom.set(
                qn("w:sz"),
                str(int(header_config.divider_thickness * 8)),
            )
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), header_config.divider_color.lstrip("#"))

            p_bdr.append(bottom)
            p_pr.append(p_bdr)

    def _render_all_sections(
        self,
        document: Document,
        cv_sections: CvSections,
        config: TemplateConfig,
    ):

        sections = cv_sections

        for section_name, data in sections.__dict__.items():
            if data is None:
                continue

            if isinstance(data, SectionMeta):
                title = data.title
                content = data.content 

                if not content:
                    continue

                section_config = config.sections.get(
                    section_name
                )

                if section_config is None:
                    section_config = SectionRendererConfig(
                        renderer=self._infer_renderer(
                            content
                        )
                    )

                self._render_section(
                    document,
                    title,
                    content,
                    section_config,
                    config,
                )

            elif isinstance(data, dict):
                for _, sub_data in data.items():
                    if not sub_data:
                        continue

                    title = sub_data.title
                    content = sub_data.content

                    section_config = SectionRendererConfig(
                        renderer=self._infer_renderer(
                            content
                        )
                    )

                    self._render_section(
                        document,
                        title,
                        content,
                        section_config,
                        config,
                    )

    def _infer_renderer(self, content) -> str:
        if isinstance(content, str):
            return "paragraph"

        if isinstance(content, list):

            if not content:
                return "paragraph"

            if isinstance(content[0], SectionEntry):
                return "entries"

            if isinstance(content[0], str):
                return "inline_list"

        raise ValueError(
            f"Cannot determine renderer for content type: "
            f"{type(content)}"
        )

    def _render_section(
        self,
        document: Document,
        title: str,
        content: str | list[SectionEntry] | list[str],
        section_config: SectionRendererConfig,
        config: TemplateConfig,
    ):

        renderer = section_config.renderer

        if renderer == "paragraph":

            self._add_str(
                document,
                title,
                content,
                config,
                section_config
            )

        elif renderer == "entries":

            self._add_bp(
                document,
                title,
                content,
                config,
                section_config,
            )

        elif renderer == "inline_list":

            self._add_list_str(
                document,
                title,
                content,
                config,
                section_config
            )

        else:

            raise ValueError(
                f"Unknown renderer: {renderer}"
            )

    def _add_str(
        self,
        document: Document,
        title: str,
        content: str,
        config: TemplateConfig,
        section_config: SectionRendererConfig,
    ):
        if not content:
            return

        if section_config.show_heading:
            self._add_heading(
                document,
                title,
                config,
                section_config,
            )

        paragraph = document.add_paragraph(content)

        space_before = (
            section_config.space_before
            if section_config.space_before is not None
            else config.section.space_before
        )

        space_after = (
            section_config.space_after
            if section_config.space_after is not None
            else config.section.space_after
        )

        paragraph.paragraph_format.space_before = Pt(space_before)
        paragraph.paragraph_format.space_after = Pt(space_after)

        run = paragraph.runs[0]
        font_config = config.font

        run.font.name = font_config.family
        run.font.size = Pt(font_config.size)
        run.bold = font_config.bold
        run.italic = font_config.italic
        run.underline = font_config.underline
        run.font.color.rgb = RGBColor.from_string(
            font_config.color.lstrip("#")
        )

    def _add_bp(
        self,
        document: Document,
        title: str,
        content: list[SectionEntry],
        config: TemplateConfig,
        section_config: SectionRendererConfig,
    ):
        if not content:
            return

        if section_config.show_heading:
            self._add_heading(
                document,
                title,
                config,
                section_config,
            )

        entry_config = (
            section_config.entry
            if section_config.entry is not None
            else config.entry
        )

        for entry in content:
            self._add_entry(
                document,
                entry,
                entry_config,
            )

    def _add_list_str(
        self,
        document: Document,
        title: str,
        content: list[str],
        config: TemplateConfig,
        section_config: SectionRendererConfig,
    ):
        if not content:
            return

        if section_config.show_heading:
            self._add_heading(
                document,
                title,
                config,
                section_config,
            )

        paragraph = document.add_paragraph()

        space_before = (
            section_config.space_before
            if section_config.space_before is not None
            else config.list.space_before
        )

        space_after = (
            section_config.space_after
            if section_config.space_after is not None
            else config.list.space_after
        )

        paragraph.paragraph_format.space_before = Pt(space_before)
        paragraph.paragraph_format.space_after = Pt(space_after)

        separator = (
            section_config.separator
            if section_config.separator is not None
            else config.list.separator
        )

        run = paragraph.add_run(separator.join(content))

        font_config = config.list.font

        run.font.name = font_config.family
        run.font.size = Pt(config.list.size)
        run.bold = font_config.bold
        run.italic = font_config.italic
        run.underline = font_config.underline
        run.font.color.rgb = RGBColor.from_string(
            font_config.color.lstrip("#")
        )

    def _add_heading(
        self,
        document: Document,
        title: str,
        config: TemplateConfig,
        section_config: SectionRendererConfig | None = None,
    ):
        if (
            section_config is not None
            and not section_config.show_heading
        ):
            return

        heading = config.section.heading

        paragraph = document.add_paragraph()

        paragraph.paragraph_format.space_before = Pt(
            heading.space_before
        )

        paragraph.paragraph_format.space_after = Pt(
            heading.space_after
        )

        if heading.uppercase:
            title = title.upper()

        run = paragraph.add_run(title)

        run.font.name = config.font.family
        run.bold = heading.bold
        run.italic = heading.italic
        run.underline = heading.underline
        run.font.size = Pt(heading.size)
        run.font.color.rgb = RGBColor.from_string(
            heading.color.lstrip("#")
        )

        if heading.show_divider:
            paragraph.paragraph_format.keep_with_next = True

            border = paragraph._p.get_or_add_pPr().get_or_add_pBdr()
            bottom = OxmlElement("w:bottom")

            bottom.set(
                qn("w:val"),
                "single",
            )
            bottom.set(
                qn("w:sz"),
                str(int(heading.divider_thickness * 8)),
            )
            bottom.set(
                qn("w:space"),
                "1",
            )
            bottom.set(
                qn("w:color"),
                heading.divider_color.lstrip("#"),
            )

            border.append(bottom)

    def _add_entry(
        self,
        document: Document,
        entry: SectionEntry,
        config: EntryConfig,
    ):
        if config.layout == "stacked":
            self._add_stacked_entry(
                document,
                entry,
                config,
            )

        elif config.layout == "compact":
            self._add_compact_entry(
                document,
                entry,
                config,
            )

        elif config.layout == "inline":
            self._add_inline_entry(
                document,
                entry,
                config,
            )

    def _add_stacked_entry(
        self,
        document: Document,
        entry: SectionEntry,
        config: EntryConfig,
    ):
        paragraph = document.add_paragraph()

        paragraph.paragraph_format.space_before = Pt(
            config.space_before
        )
        paragraph.paragraph_format.space_after = Pt(
            config.space_after
        )

        title_config = config.title

        title = paragraph.add_run(entry.title)

        self._apply_font(
            title,
            title_config,
        )

        if config.show_subtitle and entry.subtitle:
            subtitle_config = config.subtitle

            subtitle = paragraph.add_run(
                f"{config.subtitle_separator}"
                f"{entry.subtitle}"
            )

            self._apply_font(
                subtitle,
                subtitle_config,
            )

        if config.show_dates:
            self._add_dates(
                document,
                entry,
                config,
            )

        if config.show_bullets:
            self._add_bullets(
                document,
                entry.bullet_points,
                config.bullets,
            )

    def _add_compact_entry(
        self,
        document: Document,
        entry: SectionEntry,
        config: EntryConfig,
    ):
        paragraph = document.add_paragraph()

        paragraph.paragraph_format.space_before = Pt(
            config.space_before
        )
        paragraph.paragraph_format.space_after = Pt(
            config.space_after
        )

        title_config = config.title

        title = paragraph.add_run(entry.title)

        self._apply_font(
            title,
            title_config,
        )

        if config.show_subtitle and entry.subtitle:
            subtitle = paragraph.add_run(
                f"{config.subtitle_separator}"
                f"{entry.subtitle}"
            )

            self._apply_font(
                subtitle,
                config.subtitle,
            )

        if config.show_dates and self._has_dates(entry):
            dates = self._format_dates(entry, config.dates)

            date_run = paragraph.add_run(
                f"    {dates}"
            )

            self._apply_font(
                date_run,
                config.dates,
            )

        if config.show_bullets:
            self._add_bullets(
                document,
                entry.bullet_points,
                config.bullets,
            )

    def _add_inline_entry(
        self,
        document: Document,
        entry: SectionEntry,
        config: EntryConfig,
    ):
        paragraph = document.add_paragraph()

        paragraph.paragraph_format.space_before = Pt(
            config.space_before
        )
        paragraph.paragraph_format.space_after = Pt(
            config.space_after
        )

        title = paragraph.add_run(entry.title)

        self._apply_font(
            title,
            config.title,
        )

        if config.show_subtitle and entry.subtitle:
            subtitle = paragraph.add_run(
                f"{config.subtitle_separator}"
                f"{entry.subtitle}"
            )

            self._apply_font(
                subtitle,
                config.subtitle,
            )

        if config.show_dates and self._has_dates(entry):
            dates = paragraph.add_run(
                f" {config.dates.separator}"
                f"{self._format_dates(entry, config.dates)}"
            )

            self._apply_font(
                dates,
                config.dates,
            )

        if config.show_bullets:
            self._add_bullets(
                document,
                entry.bullet_points,
                config.bullets,
            )

    def _add_dates(
        self,
        document: Document,
        entry: SectionEntry,
        config: EntryConfig,
    ):
        if not self._has_dates(entry):
            return

        date_config = config.dates

        paragraph = document.add_paragraph()

        paragraph.paragraph_format.space_before = Pt(
            date_config.space_before
        )
        paragraph.paragraph_format.space_after = Pt(
            date_config.space_after
        )

        run = paragraph.add_run(
            self._format_dates(
                entry,
                date_config,
            )
        )

        self._apply_font(
            run,
            date_config,
        )

    def _add_bullets(
        self,
        document: Document,
        bullets: list[str],
        config: BulletConfig,
    ):
        for bullet in bullets:
            paragraph = document.add_paragraph()

            paragraph.paragraph_format.left_indent = Inches(
                config.indent
            )

            if config.hanging_indent is not None:
                paragraph.paragraph_format.first_line_indent = (
                    Inches(-config.hanging_indent)
                )

            paragraph.paragraph_format.space_before = Pt(
                config.space_before
            )

            paragraph.paragraph_format.space_after = Pt(
                config.space_after
            )

            paragraph.paragraph_format.line_spacing = (
                config.line_spacing
            )

            if config.alignment == "justify":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            run = paragraph.add_run(
                f"{config.symbol} {bullet}"
            )

            run.font.name = "Arial"
            run.font.size = Pt(config.size)

    def _apply_font(
        self,
        run,
        config: FontConfig,
    ):
        run.font.name = config.family
        run.font.size = Pt(config.size)
        run.bold = config.bold
        run.italic = config.italic
        run.underline = config.underline
        run.font.color.rgb = RGBColor.from_string(
            config.color.lstrip("#")
        )

    def _has_dates(
        self,
        entry: SectionEntry,
    ) -> bool:
        return bool(
            entry.start_date or entry.end_date
        )

    def _format_dates(
        self,
        entry: SectionEntry,
        config: DateConfig,
    ) -> str:
        dates = []

        if config.show_start_date and entry.start_date:
            dates.append(
                self._format_date(
                    entry.start_date,
                    config.format,
                )
            )

        if config.show_end_date and entry.end_date:
            dates.append(
                self._format_date(
                    entry.end_date,
                    config.format,
                )
            )

        if (
            config.show_end_date
            and entry.start_date
            and not entry.end_date
        ):
            dates.append(config.current_label)

        return config.separator.join(dates)

    def _format_date(
        self,
        value: str,
        format: str,
    ) -> str:
        if format == "year":
            return value[:4]

        if format == "month_year":
            if len(value) >= 7:
                return value[:7]

            return value

        if format == "full_date":
            return value

        return value

    def _convert_to_pdf(
        self,
        docx_path: Path
    ) -> str:

        output_directory = docx_path.parent

        subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(output_directory),
                str(docx_path)
            ],
            check=True
        )

        pdf_path = docx_path.with_suffix(".pdf")

        return str(pdf_path)

    def _compile_doc_url(self, path: Path) -> FileResponse:
        if not path.exists():
            raise FileNotFoundError(f"Compiled file not found: {path}")

        return FileResponse(
            path=path,
            filename=path.name,
            media_type=(
                "application/pdf"
                if path.suffix.lower() == ".pdf"
                else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ),
        )

    @staticmethod
    def to_entries(cv: CV, selected_version: str) -> list[str]:
        entries = []

        for version, sections in cv.sections.items():
            if selected_version and version != selected_version:
                continue

            for section_name, data in sections.__dict__.items():
                if data is None:
                    continue

                if isinstance(data, SectionMeta):
                    content = data.content

                    if isinstance(content, str):
                        entries.append(content)

                    elif isinstance(content, list):
                        for item in content:
                            if isinstance(item, SectionEntry):
                                entries.extend(item.bullet_points)
                            elif isinstance(item, str):
                                entries.append(item)

                elif isinstance(data, dict):
                    for _, sub_data in data.items():
                        if not sub_data:
                            continue

                        content = sub_data.content

                        if isinstance(content, str):
                            entries.append(content)

                        elif isinstance(content, list):
                            for item in content:
                                if isinstance(item, SectionEntry):
                                    entries.extend(item.bullet_points)
                                elif isinstance(item, str):
                                    entries.append(item)

        return entries

    @staticmethod
    def to_json(sc: CV) -> dict:
        return sc.model_dump()
        
    @staticmethod
    def from_json(json_data: dict) -> CV:
        return CV.model_validate(json_data)
    
    @staticmethod
    def _safe_filename(
        name: str
    ) -> str:

        return "".join(
            c for c in name
            if c.isalnum() or c in " _-"
        ).strip()