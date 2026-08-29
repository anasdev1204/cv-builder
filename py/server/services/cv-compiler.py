from pathlib import Path
from tempfile import TemporaryDirectory
import subprocess

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

from models.cv import CV, SectionEntry


class CVCompiler:

    def __init__(
        self,
        output_directory: str = "generated"
    ):
        self.output_directory = Path(output_directory)
        self.output_directory.mkdir(
            parents=True,
            exist_ok=True
        )

    async def compile(
        self,
        cv: CV,
        template: str = "professional",
        output_format: str = "pdf"
    ) -> str:

        docx_path = self._compile_docx(
            cv=cv,
            template=template
        )

        if output_format == "docx":
            return str(docx_path)

        if output_format == "pdf":
            return self._convert_to_pdf(docx_path)

        raise ValueError(
            f"Unsupported format: {output_format}"
        )

    def _compile_docx(
        self,
        cv: CV,
        template: str
    ) -> Path:

        document = Document()

        self._configure_page(document)
        self._configure_styles(document)

        self._add_header(
            document,
            cv
        )

        self._add_summary(
            document,
            cv
        )

        self._add_experience(
            document,
            cv
        )

        self._add_education(
            document,
            cv
        )

        self._add_skills(
            document,
            cv
        )

        self._add_languages(
            document,
            cv
        )

        self._add_other_sections(
            document,
            cv
        )

        output_path = (
            self.output_directory /
            f"{self._safe_filename(cv.userData.name)}.docx"
        )

        document.save(output_path)

        return output_path

    def _configure_page(
        self,
        document: Document
    ):

        section = document.sections[0]

        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)

    def _configure_styles(
        self,
        document: Document
    ):

        styles = document.styles

        normal = styles["Normal"]

        normal.font.name = "Arial"
        normal.font.size = Pt(9.5)

        normal.paragraph_format.space_after = Pt(2)

    def _add_header(
        self,
        document: Document,
        cv: CV
    ):

        user = cv.userData

        paragraph = document.add_paragraph()

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        name = paragraph.add_run(user.name)

        name.bold = True
        name.font.size = Pt(20)

        contact = []

        if user.email:
            contact.append(user.email)

        if user.phoneNumber:
            contact.append(user.phoneNumber)

        if user.address:
            location = (
                f"{user.address.city}, "
                f"{user.address.country}"
            )

            contact.append(location)

        if user.linkedin:
            contact.append(user.linkedin)

        if user.portfolio:
            contact.append(user.portfolio)

        if contact:

            paragraph = document.add_paragraph()

            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run = paragraph.add_run(
                " | ".join(contact)
            )

            run.font.size = Pt(8.5)

    def _add_summary(
        self,
        document: Document,
        cv: CV
    ):

        if not cv.sections.summary:
            return

        self._add_section_heading(
            document,
            "SUMMARY"
        )

        paragraph = document.add_paragraph(
            cv.sections.summary
        )

        paragraph.paragraph_format.space_after = Pt(5)

    def _add_experience(
        self,
        document: Document,
        cv: CV
    ):

        if not cv.sections.experience:
            return

        self._add_section_heading(
            document,
            "EXPERIENCE"
        )

        for entry in cv.sections.experience:

            self._add_entry(
                document,
                entry
            )

    def _add_education(
        self,
        document: Document,
        cv: CV
    ):

        if not cv.sections.education:
            return

        self._add_section_heading(
            document,
            "EDUCATION"
        )

        for entry in cv.sections.education:

            self._add_entry(
                document,
                entry
            )

    def _add_skills(
        self,
        document: Document,
        cv: CV
    ):

        if not cv.sections.skills:
            return

        self._add_section_heading(
            document,
            "SKILLS"
        )

        paragraph = document.add_paragraph()

        paragraph.add_run(
            " • ".join(cv.sections.skills)
        )

    def _add_languages(
        self,
        document: Document,
        cv: CV
    ):

        if not cv.sections.languages:
            return

        self._add_section_heading(
            document,
            "LANGUAGES"
        )

        paragraph = document.add_paragraph()

        paragraph.add_run(
            " • ".join(cv.sections.languages)
        )

    def _add_other_sections(
        self,
        document: Document,
        cv: CV
    ):

        for section_name, entries in (
            cv.sections.otherSections.items()
        ):

            if not entries:
                continue

            self._add_section_heading(
                document,
                section_name.upper()
            )

            for entry in entries:

                self._add_entry(
                    document,
                    entry
                )

    def _add_section_heading(
        self,
        document: Document,
        title: str
    ):

        paragraph = document.add_paragraph()

        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(3)

        run = paragraph.add_run(title)

        run.bold = True
        run.font.size = Pt(10.5)

    def _add_entry(
        self,
        document: Document,
        entry: SectionEntry
    ):

        paragraph = document.add_paragraph()

        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(0)

        title = paragraph.add_run(
            entry.title
        )

        title.bold = True
        title.font.size = Pt(10)

        if entry.subtitle:

            subtitle = paragraph.add_run(
                f" — {entry.subtitle}"
            )

            subtitle.font.size = Pt(9.5)

        dates = []

        if entry.startDate:
            dates.append(entry.startDate)

        if entry.endDate:
            dates.append(entry.endDate)

        if dates:

            date_paragraph = document.add_paragraph()

            date_paragraph.paragraph_format.space_after = Pt(1)

            run = date_paragraph.add_run(
                " – ".join(dates)
            )

            run.italic = True
            run.font.size = Pt(8.5)

        for bullet in entry.bulletPoints:

            paragraph = document.add_paragraph(
                style="List Bullet"
            )

            paragraph.paragraph_format.left_indent = Inches(
                0.2
            )

            paragraph.paragraph_format.space_after = Pt(1)

            run = paragraph.add_run(bullet)

            run.font.size = Pt(9)

    def _convert_to_pdf(
        self,
        docx_path: Path
    ) -> str:

        output_directory = docx_path.parent

        subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(output_directory),
                str(docx_path)
            ],
            check=True
        )

        pdf_path = docx_path.with_suffix(".pdf")

        return str(pdf_path)

    @staticmethod
    def _safe_filename(
        name: str
    ) -> str:

        return "".join(
            c for c in name
            if c.isalnum() or c in " _-"
        ).strip()